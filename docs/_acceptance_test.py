import sys, os, tempfile
from PySide6.QtWidgets import QApplication, QFileDialog
from PySide6.QtCore import QTimer
import src.main_window as mw_mod
from src.main_window import MainWindow
from src.file_entry import FileStatus

d = tempfile.mkdtemp(prefix="acc_")
# --- generate real test files ---
open(d+"/s.txt","w",encoding="utf-8").write("# Hi\n\n*bold*")
from fpdf import FPDF
p=FPDF(); p.add_page(); p.set_font("Helvetica",size=20); p.cell(0,10,"Acc PDF"); p.output(d+"/s.pdf")
from docx import Document
doc=Document(); doc.add_heading("Acc Doc",0); doc.add_paragraph("body"); doc.save(d+"/s.docx")
from openpyxl import Workbook
wb=Workbook(); ws=wb.active; ws["A1"]="X"; ws["B1"]="Y"; ws.append([1,2]); wb.save(d+"/s.xlsx")
# error files
open(d+"/corrupt.pdf","wb").write(b"%PDF-1.4\n%%EOF\nthis is not a real pdf garbage")
open(d+"/bad.bin","wb").write(os.urandom(120))   # binary -> unsupported, must not crash

PDF=d+"/s.pdf"; DOCX=d+"/s.docx"; XLSX=d+"/s.xlsx"; CORRUPT=d+"/corrupt.pdf"; BAD=d+"/bad.bin"

app = QApplication(sys.argv)
app.setApplicationName("acceptance")
res={}
def ok(m): res.setdefault("OK",[]).append(m); print("OK  ",m)
def fail(m): res.setdefault("FAIL",[]).append(m); print("FAIL",m)

w = MainWindow()
w.show()

SAVE=os.path.join(tempfile.gettempdir(),"acc_export.md")
def fake_save(*a,**k): return (SAVE,"Markdown (*.md)")
def fake_cancel(*a,**k): return ("","Markdown (*.md)")
QFileDialog.getSaveFileName = fake_save

state={"s":0}; markdown_of_done=None; done_row=None

def step():
    try:
        if state["s"]==0:
            w._model.add_paths([PDF,DOCX,XLSX,CORRUPT,BAD])
            w._model.add_paths([PDF])  # duplicate -> dedup
            if w._model.rowCount()==5: ok("选择/拖入去重: 5 个条目")
            else: fail("条目数应为5, 实为 %d"%w._model.rowCount())
            w.start_conversion()
            state["s"]=1
        elif state["s"]==1:
            if not w._converting:
                # result assertions
                NP=os.path.normcase
                for r in range(w._model.rowCount()):
                    e=w._model.entry_at(r)
                    if NP(e.path)==NP(PDF):
                        if e.status==FileStatus.DONE and e.markdown: ok("PDF 转换=DONE")
                        else: fail("PDF 状态=%s"%e.status)
                    if NP(e.path)==NP(DOCX):
                        if e.status==FileStatus.DONE and e.markdown: ok("DOCX 转换=DONE")
                        else: fail("DOCX 状态=%s"%e.status)
                    if NP(e.path)==NP(XLSX):
                        if e.status==FileStatus.DONE and e.markdown: ok("XLSX 转换=DONE")
                        else: fail("XLSX 状态=%s"%e.status)
                    if NP(e.path)==NP(CORRUPT):
                        if e.status==FileStatus.ERROR and e.error_message: ok("损坏PDF=ERROR(不崩溃)")
                        else: fail("损坏PDF 状态=%s msg=%s"%(e.status,e.error_message))
                    if NP(e.path)==NP(BAD):
                        if e.status in (FileStatus.ERROR,FileStatus.UNSUPPORTED): ok("二进制错误文件=%s(不崩溃)"%e.status)
                        else: fail("二进制文件状态=%s"%e.status)
                # find a DONE row for preview/copy/export
                for r in range(w._model.rowCount()):
                    e=w._model.entry_at(r)
                    if e.status==FileStatus.DONE and e.markdown:
                        done_row=r; markdown_of_done=e.markdown; break
                state["s"]=2
                w._current_row=done_row; w._show_entry(done_row)
                # source + preview
                if w._source_view.toPlainText().strip(): ok("Markdown 源码查看正常")
                else: fail("源码视图为空")
                if w._preview_view.toHtml().strip(): ok("渲染预览正常(toHtml 非空)")
                else: fail("预览为空")
                if w._act_copy.isEnabled() and w._act_export.isEnabled(): ok("复制/导出按钮随 DONE 自动启用")
                else: fail("DONE 下复制/导出按钮未启用")
                # copy
                try:
                    w._on_copy(); ok("复制: 无异常")
                except Exception as ex: fail("复制抛异常: %r"%ex)
                # export happy
                try:
                    if os.path.exists(SAVE): os.remove(SAVE)
                    w._on_export()
                    if os.path.exists(SAVE) and open(SAVE,encoding="utf-8").read()==markdown_of_done:
                        ok("导出 .md: 内容正确(UTF-8)")
                    else: fail("导出文件缺失或内容不符")
                except Exception as ex: fail("导出抛异常: %r"%ex)
                # export cancel
                QFileDialog.getSaveFileName=fake_cancel
                try:
                    w._on_export(); ok("导出取消: 无异常")
                except Exception as ex: fail("导出取消抛异常: %r"%ex)
                QFileDialog.getSaveFileName=fake_save
                state["s"]=3
                w.close()
                w2=MainWindow(); w2.show(); ok("关闭后可再次启动(relaunch)")
                w2.close()
                finish()
            # else: still converting, keep polling
        else:
            pass
    except Exception as ex:
        fail("未捕获异常: %r"%ex); finish()

def finish():
    print("\n===== ACCEPTANCE SUMMARY =====")
    print("OK=%d  FAIL=%d"%(len(res.get("OK",[])),len(res.get("FAIL",[]))))
    if res.get("FAIL"):
        for f in res["FAIL"]: print("  -",f)
        app.quit(); sys.exit(1)
    else:
        print("ALL ACCEPTANCE CHECKS PASSED")
        app.quit(); sys.exit(0)

t=QTimer(); t.timeout.connect(step); t.start(150)
QTimer.singleShot(40000, lambda: (fail("超时未完成(step %d)"%state["s"]), finish()))
app.exec()
