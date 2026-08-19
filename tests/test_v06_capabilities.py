"""MdDesk v0.6 capability decoupling & error isolation tests (plan §3.2/§3.4).

Verifies, at THREE levels:

  A. Factory wiring (spy level): the four OCR / image-description toggle
     combinations produce the correct engine wiring —
       * both off        -> plain MarkItDown(), no llm kwargs, no plugin
       * OCR only        -> plugin registered WITH its own llm kwargs,
                            constructor gets none (no image description)
       * desc only       -> constructor llm kwargs, plugin NOT registered
       * both            -> both wirings
     plus: AI disabled -> plain MarkItDown() (v0.2 path); client built once
     and shared (cached).

  B. Real conversion (mock LLM client): a PNG + a DOCX-with-image convert
     under each combination and the markdown reflects exactly the enabled
     capabilities (description block / OCR block present or absent).

  C. Worker -> result -> report chain: an AI provider failure is isolated —
     the file still converts (DONE), the result carries an
     AI_PROVIDER_FAILURE warning, the ConversionReport + diagnostic log
     contain the sanitized warning and NEVER the API key.

Run: python tests/test_v06_capabilities.py  (or pytest tests/test_v06_capabilities.py)
Headless (offscreen Qt only for the worker part), network-free.
"""

import io
import json
import os
import sys
import tempfile
from pathlib import Path
from unittest.mock import patch

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

import openai  # noqa: E402
from markitdown import MarkItDown  # noqa: E402

from src.converter import convert_file  # noqa: E402
from src.engine_config import EngineConfig  # noqa: E402
from src.file_entry import FileEntry, FileStatus  # noqa: E402
from src.markitdown_factory import MarkItDownFactory  # noqa: E402
from src.report import DiagnosticLogger, ReportBuilder  # noqa: E402
from src.result import ConversionResult  # noqa: E402
from src.settings import Settings, AIConfig  # noqa: E402
from src.worker import ConversionWorker  # noqa: E402

SECRET = "sk-SECRET-FOR-ISOLATION-TEST"


def _check(name, cond, detail=""):
    print(("PASS" if cond else "FAIL"), "-", name, "" if cond else f":: {detail}")
    return cond


# --------------------------------------------------------------------------- #
# shared fixtures                                                             #
# --------------------------------------------------------------------------- #
class _Msg:
    def __init__(self, c):
        self.content = c


class _Choice:
    def __init__(self, c):
        self.message = _Msg(c)


class _Resp:
    def __init__(self, c):
        self.choices = [_Choice(c)]


class MockClient:
    """Canned LLM: returns 'MOCK_AI_TEXT' or raises the chosen exception."""

    def __init__(self, exc=None):
        self.exc = exc
        self.calls = 0
        client = self

        class _Completions:
            def create(self, **kw):
                client.calls += 1
                if client.exc is not None:
                    raise client.exc
                return _Resp("MOCK_AI_TEXT")

        class _Chat:
            completions = _Completions()

        self.chat = _Chat()


def _fake_resp(code):
    class R:
        status_code = code
        headers = {}
        request = None
        is_success = False
    return R()


def _mk_err(kind):
    if kind == "auth":
        return openai.AuthenticationError(message="bad key", response=_fake_resp(401), body=None)
    if kind == "timeout":
        return openai.APITimeoutError(request=None)
    raise ValueError(kind)


def _make_png(path: Path) -> None:
    from PIL import Image
    Image.new("RGB", (64, 64), (10, 10, 200)).save(path, format="PNG")


def _make_docx_with_image(path: Path) -> None:
    from docx import Document
    import io as _io
    from PIL import Image
    buf = _io.BytesIO()
    Image.new("RGB", (32, 32), (200, 10, 10)).save(buf, format="PNG")
    d = Document()
    d.add_paragraph("Hello Docx")
    d.add_picture(_io.BytesIO(buf.getvalue()))
    d.save(str(path))


def _entry(path: Path) -> FileEntry:
    return FileEntry(path=str(path), filename=path.name,
                     extension=path.suffix.lower(), size=path.stat().st_size)


def _cfg(ocr: bool, desc: bool, client) -> EngineConfig:
    cfg = EngineConfig(
        ai_enabled=True, ai_endpoint="", ai_api_key=SECRET,
        ai_model="gpt-4o", ai_ocr_enabled=ocr, ai_image_description_enabled=desc,
    )
    cfg._client = client
    return cfg


# --------------------------------------------------------------------------- #
# A. Factory wiring for the four combinations                                 #
# --------------------------------------------------------------------------- #
def test_factory_wiring_matrix():
    ok = True
    import markitdown_ocr

    class _FakeMD:
        instances = []

        def __init__(self, enable_plugins=None, **kwargs):
            self.kwargs = kwargs
            self.enable_plugins = enable_plugins
            self.registered = []
            _FakeMD.instances.append(self)

    for ocr, desc in ((False, False), (False, True), (True, False), (True, True)):
        _FakeMD.instances.clear()
        client = MockClient()
        with patch("src.markitdown_factory.MarkItDown", _FakeMD), \
             patch.object(markitdown_ocr, "register_converters",
                          lambda md, **kw: md.registered.append(kw)):
            md = MarkItDownFactory.create(_cfg(ocr, desc, client))
        tag = f"ocr={int(ocr)},desc={int(desc)}"
        ok &= _check(f"A.{tag} 引擎创建", md is not None)
        ok &= _check(f"A.{tag} 构造 kwargs 含 llm_client = desc 开",
                     ("llm_client" in md.kwargs) is desc, md.kwargs.keys())
        ok &= _check(f"A.{tag} OCR 插件注册 = ocr 开",
                     bool(md.registered) is ocr)
        if ocr:
            ok &= _check(f"A.{tag} 插件拿到独立 llm kwargs",
                         md.registered[0].get("llm_client") is client)

    # AI disabled -> plain engine, no kwargs, no plugin, NO client built.
    _FakeMD.instances.clear()
    with patch("src.markitdown_factory.MarkItDown", _FakeMD), \
         patch.object(markitdown_ocr, "register_converters",
                      lambda md, **kw: md.registered.append(kw)):
        md = MarkItDownFactory.create(EngineConfig())
    ok &= _check("A.off AI 关闭 -> 无 llm kwargs", "llm_client" not in md.kwargs)
    ok &= _check("A.off AI 关闭 -> 不注册插件", not md.registered)

    # Master on, both capabilities off -> plain engine too (nothing AI-shaped).
    _FakeMD.instances.clear()
    with patch("src.markitdown_factory.MarkItDown", _FakeMD), \
         patch.object(markitdown_ocr, "register_converters",
                      lambda md, **kw: md.registered.append(kw)):
        md = MarkItDownFactory.create(_cfg(False, False, MockClient()))
    ok &= _check("A.both-off -> 无 llm kwargs", "llm_client" not in md.kwargs)
    ok &= _check("A.both-off -> 不注册插件", not md.registered)

    # Client shared across engines built from the same config (batch reuse).
    _FakeMD.instances.clear()
    cfg = _cfg(True, True, MockClient())
    with patch("src.markitdown_factory.MarkItDown", _FakeMD), \
         patch.object(markitdown_ocr, "register_converters",
                      lambda md, **kw: md.registered.append(kw)):
        md1 = MarkItDownFactory.create(cfg)
        md2 = MarkItDownFactory.create(cfg)
    ok &= _check("A.批量共享同一 client",
                 md1.kwargs["llm_client"] is md2.kwargs["llm_client"])
    return ok


# --------------------------------------------------------------------------- #
# B. Real conversion under the four combinations                              #
# --------------------------------------------------------------------------- #
def test_conversion_capability_matrix():
    ok = True
    with tempfile.TemporaryDirectory() as d:
        png = Path(d) / "img.png"
        docx = Path(d) / "doc.docx"
        _make_png(png)
        _make_docx_with_image(docx)

        for ocr, desc in ((False, False), (False, True), (True, False), (True, True)):
            tag = f"ocr={int(ocr)},desc={int(desc)}"
            client = MockClient()
            md_png = convert_file(str(png), engine_config=_cfg(ocr, desc, client))
            md_docx = convert_file(str(docx), engine_config=_cfg(ocr, desc, client))

            ok &= _check(f"B.{tag} PNG 描述块 = desc 开",
                         ("# Description:" in md_png) is desc,
                         md_png[:80])
            ok &= _check(f"B.{tag} DOCX OCR 块 = ocr 开",
                         ("[Image OCR]" in md_docx) is ocr,
                         md_docx[:120])
            ok &= _check(f"B.{tag} 转换均为成功产出",
                         isinstance(md_png, str) and isinstance(md_docx, str)
                         and len(md_docx) > 0)
    return ok


# --------------------------------------------------------------------------- #
# C. Provider failure isolation through the worker -> report chain            #
# --------------------------------------------------------------------------- #
def _run_worker(tasks, engine_config, quality_enabled=False):
    """Run a ConversionWorker synchronously and collect its results."""
    from PySide6.QtWidgets import QApplication
    from PySide6.QtCore import QEventLoop, QTimer
    app = QApplication.instance() or QApplication([])
    worker = ConversionWorker(tasks, engine_config=engine_config,
                              quality_enabled=quality_enabled)
    results = []
    worker.file_finished.connect(results.append)
    loop = QEventLoop()
    worker.batch_finished.connect(lambda *a: loop.quit())
    worker.batch_cancelled.connect(lambda *a: loop.quit())
    worker.start()
    QTimer.singleShot(120000, loop.quit)
    loop.exec()
    while not worker.isFinished():
        app.processEvents()
    return results


def test_provider_failure_isolation_and_report():
    ok = True
    with tempfile.TemporaryDirectory() as d:
        png = Path(d) / "img.png"
        _make_png(png)
        entry = _entry(png)
        tasks = [(0, entry)]

        # AI on, provider raises auth (secret-bearing) error.
        cfg = _cfg(True, True, MockClient(exc=_mk_err("auth")))
        results = _run_worker(tasks, cfg)
        ok &= _check("C1. worker 产出结果", len(results) == 1)

        r = results[0]
        ok &= _check("C2. AI 故障被隔离：文件仍 DONE", r.status == FileStatus.DONE,
                     f"status={r.status}")
        ok &= _check("C3. 结果携带 AI_PROVIDER_FAILURE warning",
                     any(w.code == "AI_PROVIDER_FAILURE" for w in r.warnings),
                     [w.code for w in r.warnings])
        warn = next(w for w in r.warnings if w.code == "AI_PROVIDER_FAILURE")
        ok &= _check("C4. warning 含 provider/model 维度",
                     "provider=openai-compatible" in warn.message
                     and "model=gpt-4o" in warn.message, warn.message[:100])
        ok &= _check("C5. markdown 为无 AI 降级产物",
                     "Description:" not in (r.markdown or ""))

        # Report + diagnostic log: sanitized, no secret.
        report = ReportBuilder.build(r, entry)
        ok &= _check("C6. report 状态 DONE", report.ok)
        ok &= _check("C7. report warnings 含 AI 故障",
                     any(w.code == "AI_PROVIDER_FAILURE" for w in report.warnings))
        log = DiagnosticLogger(log_dir=d)
        ok &= _check("C8. 日志写入成功", log.record(report) is True)
        line = Path(log.log_path).read_text(encoding="utf-8")
        parsed = json.loads(line.strip())
        ok &= _check("C9. 日志含 AI 故障 code",
                     any(w["code"] == "AI_PROVIDER_FAILURE" for w in parsed["warnings"]))
        blob = line + json.dumps(report.to_dict()) + warn.message
        ok &= _check("C10. 日志/report/warning 不含 API Key", SECRET not in blob)

        # Timeout variant: provider times out -> same isolation.
        cfg_t = _cfg(True, True, MockClient(exc=_mk_err("timeout")))
        results_t = _run_worker(tasks, cfg_t)
        r_t = results_t[0]
        ok &= _check("C11. 超时同样隔离（DONE）", r_t.status == FileStatus.DONE)
        ok &= _check("C12. 超时 warning 分类正确",
                     any("超时" in w.message for w in r_t.warnings if w.code == "AI_PROVIDER_FAILURE"))
    return ok


def test_no_ai_config_normal_conversion():
    """The most important default user path: no AI config at all."""
    ok = True
    with tempfile.TemporaryDirectory() as d:
        png = Path(d) / "img.png"
        _make_png(png)
        settings = Settings.default()  # AI off
        ok &= _check("N1. 默认设置 AI 关闭", not settings.ai.enabled)

        cfg = EngineConfig.from_settings(settings)
        ok &= _check("N2. from_settings -> disabled", not cfg.ai_enabled)

        md = convert_file(str(png), engine_config=cfg)
        ok &= _check("N3. 普通转换成功", isinstance(md, str))
        ok &= _check("N4. 无 AI 内容", "Description:" not in md and "MOCK" not in md)
        entry = _entry(png)
        results = _run_worker([(0, entry)], cfg)
        ok &= _check("N5. worker 普通路径 DONE", results and results[0].ok)
        ok &= _check("N6. 无 warning", results and not results[0].warnings)
    return ok


def _main():
    ok = True
    ok &= test_factory_wiring_matrix()
    ok &= test_conversion_capability_matrix()
    ok &= test_provider_failure_isolation_and_report()
    ok &= test_no_ai_config_normal_conversion()
    print("RESULT:", "PASS" if ok else "FAIL")
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(_main())
