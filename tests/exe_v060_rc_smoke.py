"""Frozen-binary RC smoke for MdDesk v0.6.0 (AI 实用化版).

Headless (offscreen). Drives the real MainWindow + worker + v0.6 AI provider
infrastructure inside the PyInstaller-frozen EXE (mirrors md-desk.spec's
collection settings). Also runnable as a source script for pre-build
validation (frozen-only checks degrade gracefully).

Coverage (all offline-safe EXCEPT two localhost probes — no external network):

  * 无 AI 配置启动（最重要默认路径）: APPDATA redirected to a temp dir ->
    default settings, AI off, MainWindow boots, txt converts DONE with no
    warnings
  * v0.5 settings.json migration inside the frozen app: old ai block loads
    with the v0.6 defaults (provider/timeout/ocr/description)
  * ai_provider layer works frozen: ClientFactory builds a client with a
    finite timeout and max_retries=0
  * Connection test (REAL localhost round-trips):
      - a local OpenAI-compatible mock endpoint -> SUCCESS
      - an endpoint returning 401 -> auth failure classified
      - a silent (accept-but-hang) endpoint -> TIMEOUT classified
  * Capability toggles (mock LLM client injected): PNG description on/off,
    DOCX embedded-image OCR on/off — the explicit plugin registration path
    used identically in dev and frozen
  * AI failure isolation: mock auth error -> file still DONE with an
    AI_PROVIDER_FAILURE warning; report/log carry no API key
  * GUI regression: local file real MarkItDown conversion DONE

Run both as a source script and as the frozen RC smoke EXE.
"""

import io
import json
import os
import socket
import sys
import tempfile
import threading
import time
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

# ---- isolate %APPDATA% BEFORE importing anything from src ----------------
# Guarantees the frozen EXE starts with NO prior settings (the default user
# path) and writes logs into the temp dir instead of the user's real one.
_SMOKE_APPDATA = tempfile.mkdtemp(prefix="mdesk_v060_smoke_")
os.environ["APPDATA"] = _SMOKE_APPDATA

ROOT = Path(__file__).resolve().parent.parent
if not getattr(sys, "frozen", False):
    sys.path.insert(0, str(ROOT))

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from PySide6.QtCore import QEventLoop, QTimer  # noqa: E402
from PySide6.QtWidgets import QApplication  # noqa: E402

from src import ai_provider  # noqa: E402
from src.ai_provider import AIProviderConfig, ClientFactory  # noqa: E402
from src.converter import convert_file  # noqa: E402
from src.engine_config import EngineConfig  # noqa: E402
from src.file_entry import FileEntry, FileStatus  # noqa: E402
from src.markitdown_factory import MarkItDownFactory  # noqa: E402
from src.report import DiagnosticLogger, ReportBuilder  # noqa: E402
from src.settings import AIConfig, Settings  # noqa: E402
from src.worker import ConversionWorker  # noqa: E402

_OUT = []
SECRET = "sk-FROZEN-SMOKE-SECRET-2468"


def _log(name, ok, detail=""):
    _OUT.append((name, ok))
    print(("PASS" if ok else "FAIL"), "-", name, "" if ok else f":: {detail}", flush=True)
    return ok


def _wait_batch_done(window, timeout_ms=60000):
    app = QApplication.instance()
    loop = QEventLoop()

    def _poll():
        if window._worker is None:
            loop.quit()

    timer = QTimer()
    timer.timeout.connect(_poll)
    timer.start(50)
    QTimer.singleShot(timeout_ms, loop.quit)
    loop.exec()
    timer.stop()
    app.processEvents()
    app.sendPostedEvents()
    return window._worker is None


# ---------------------------------------------------------------------------
# Local OpenAI-compatible mock endpoints (localhost only)
# ---------------------------------------------------------------------------
class _Handler(BaseHTTPRequestHandler):
    def log_message(self, *args):  # silence
        pass

    def do_POST(self):  # noqa: N802
        path = self.path.lower()
        if path.startswith("/auth"):
            body = json.dumps({"error": {"message": "bad key",
                                         "type": "invalid_request_error"}}).encode()
            self.send_response(401)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return
        if path.startswith("/slow"):
            time.sleep(30)  # accept but never answer in time
            self.send_response(200)
            self.end_headers()
            return
        # /ok... -> a minimal valid chat.completions response
        body = json.dumps({
            "id": "chatcmpl-smoke", "object": "chat.completion",
            "created": 1, "model": "gpt-4o",
            "choices": [{"index": 0, "finish_reason": "stop",
                         "message": {"role": "assistant", "content": "pong"}}],
        }).encode()
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def _start_mock_server():
    srv = ThreadingHTTPServer(("127.0.0.1", 0), _Handler)
    threading.Thread(target=srv.serve_forever, daemon=True).start()
    return srv


# ---------------------------------------------------------------------------
# Mock LLM client (capability toggles + isolation; no network)
# ---------------------------------------------------------------------------
class _Msg:
    def __init__(self, c):
        self.content = c


class _Choice:
    def __init__(self, c):
        self.message = _Msg(c)


class _Resp:
    def __init__(self, c):
        self.choices = [_Choice(c)]


class MockClient:
    def __init__(self, exc=None):
        self.exc = exc
        self.calls = 0
        client = self

        class _Completions:
            def create(self, **kw):
                client.calls += 1
                if client.exc is not None:
                    raise client.exc
                return _Resp("MOCK_AI_TEXT")

        class _Chat:
            completions = _Completions()

        self.chat = _Chat()


def _mk_auth_error():
    import openai

    class R:
        status_code = 401
        headers = {}
        request = None
        is_success = False

    return openai.AuthenticationError(message="bad key " + SECRET,
                                      response=R(), body=None)


def _make_png(path: Path) -> None:
    from PIL import Image
    Image.new("RGB", (64, 64), (10, 10, 200)).save(path, format="PNG")


def _make_docx_with_image(path: Path) -> None:
    from docx import Document
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", (32, 32), (200, 10, 10)).save(buf, format="PNG")
    d = Document()
    d.add_paragraph("Hello Docx")
    d.add_picture(io.BytesIO(buf.getvalue()))
    d.save(str(path))


def _cfg(ocr: bool, desc: bool, client) -> EngineConfig:
    cfg = EngineConfig(
        ai_enabled=True, ai_endpoint="", ai_api_key=SECRET,
        ai_model="gpt-4o", ai_ocr_enabled=ocr,
        ai_image_description_enabled=desc,
    )
    cfg._client = client
    return cfg


# ---------------------------------------------------------------------------
# 1. No-AI-config boot (THE default user path)
# ---------------------------------------------------------------------------
def _no_ai_boot_check(window):
    s = Settings.load()
    ok = _log("V6.1 无 AI 配置: 默认 settings 加载且 AI 关",
              s.ai.enabled is False and not s.ai.model)
    ok &= _log("V6.1b 默认 timeout 有限",
               1.0 <= s.ai.timeout_seconds <= 600.0, s.ai.timeout_seconds)
    d = tempfile.mkdtemp()
    p = Path(d) / "note.txt"
    p.write_text("# Title\n\nSome real content.\n", encoding="utf-8")
    window._model.add_paths([str(p)])
    window.start_conversion()
    if not _wait_batch_done(window):
        return _log("V6.1 批次结束", False, "timeout")
    entry = window._model.entry_at(0)
    ok &= _log("V6.1c 无 AI 配置转换 DONE 无警告",
               entry.status == FileStatus.DONE and bool(entry.markdown)
               and entry.report is not None and entry.report.warnings == (),
               f"{entry.status.value}")
    window._model.clear()
    return ok


# ---------------------------------------------------------------------------
# 2. v0.5 settings migration inside the frozen app
# ---------------------------------------------------------------------------
def _migration_check():
    p = Path(_SMOKE_APPDATA) / "MdDesk" / "settings.json"
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps({
        "version": 2,
        "youtube_transcript_languages": ["zh-Hans"],
        "ai": {"enabled": True, "endpoint": "https://api.openai.com/v1",
               "model": "gpt-4o", "prompt": ""},
        "quality_enabled": True,
    }), encoding="utf-8")
    s = Settings.load(p)
    ok = _log("V6.2 v0.5 设置迁移: enabled/model 保留",
              s.ai.enabled and s.ai.model == "gpt-4o")
    ok &= _log("V6.2b 迁移默认: provider/timeout/两开关",
               s.ai.provider == "openai-compatible"
               and s.ai.timeout_seconds == 60.0
               and s.ai.ocr_enabled and s.ai.image_description_enabled)
    # Clean up so later checks keep the "no settings" state.
    p.unlink()
    return ok


# ---------------------------------------------------------------------------
# 3. Provider layer works frozen
# ---------------------------------------------------------------------------
def _provider_layer_check():
    ok = True
    client = ClientFactory.create(AIProviderConfig(
        api_key=SECRET, endpoint="https://gw.example/v1",
        model="gpt-4o", timeout_seconds=15))
    ok = _log("V6.3 ClientFactory 构建客户端", client is not None)
    ok &= _log("V6.3b timeout=15 传入", float(client.timeout) == 15.0, client.timeout)
    ok &= _log("V6.3c max_retries=0", getattr(client, "max_retries", -1) == 0)
    ok &= _log("V6.3d describe 不含 key",
               SECRET not in ai_provider.AIProviderConfig(
                   api_key=SECRET, endpoint="https://x/v1").describe())
    return ok


# ---------------------------------------------------------------------------
# 4. Connection test — REAL localhost round-trips
# ---------------------------------------------------------------------------
def _connection_test_check(srv):
    port = srv.server_address[1]
    base = f"http://127.0.0.1:{port}"
    ok = True

    # success
    r = ai_provider.test_connection(AIProviderConfig(
        api_key=SECRET, endpoint=base + "/ok/v1", model="gpt-4o",
        timeout_seconds=10))
    ok &= _log("V6.4 连接测试 success（本地真实 HTTP）", r.ok, r.message)
    ok &= _log("V6.4b success 消息不含 key", SECRET not in r.message)

    # 401 -> auth failure
    r = ai_provider.test_connection(AIProviderConfig(
        api_key=SECRET, endpoint=base + "/auth/v1", model="gpt-4o",
        timeout_seconds=10))
    ok &= _log("V6.4c 401 -> 鉴权失败分类", (not r.ok) and "401" in r.message, r.message)
    ok &= _log("V6.4d 鉴权失败消息不含 key", SECRET not in r.message)

    # silent endpoint -> timeout
    r = ai_provider.test_connection(AIProviderConfig(
        api_key=SECRET, endpoint=base + "/slow/v1", model="gpt-4o",
        timeout_seconds=2))
    ok &= _log("V6.4e 无响应端点 -> 超时分类", (not r.ok) and "超时" in r.message, r.message)

    # unreachable port -> connection failure
    r = ai_provider.test_connection(AIProviderConfig(
        api_key=SECRET, endpoint="http://127.0.0.1:1/v1", model="gpt-4o",
        timeout_seconds=5))
    ok &= _log("V6.4f 不可达端口 -> 连接失败分类",
               (not r.ok) and ("无法连接" in r.message or "超时" in r.message),
               r.message)
    return ok


# ---------------------------------------------------------------------------
# 5. Capability toggles (mock client) — explicit plugin registration path
# ---------------------------------------------------------------------------
def _capability_toggle_check():
    ok = True
    d = tempfile.mkdtemp()
    png = Path(d) / "img.png"
    docx = Path(d) / "doc.docx"
    _make_png(png)
    _make_docx_with_image(docx)

    for ocr, desc in ((False, False), (False, True), (True, False), (True, True)):
        tag = f"ocr={int(ocr)},desc={int(desc)}"
        client = MockClient()
        md_png = convert_file(str(png), engine_config=_cfg(ocr, desc, client))
        md_docx = convert_file(str(docx), engine_config=_cfg(ocr, desc, client))
        ok &= _log(f"V6.5 {tag} PNG 描述块=desc开",
                   ("# Description:" in md_png) is desc, md_png[:60])
        ok &= _log(f"V6.5b {tag} DOCX OCR块=ocr开",
                   ("[Image OCR]" in md_docx) is ocr, md_docx[:80])
    return ok


# ---------------------------------------------------------------------------
# 6. AI failure isolation through the real worker
# ---------------------------------------------------------------------------
def _run_worker(tasks, engine_config):
    app = QApplication.instance() or QApplication([])
    worker = ConversionWorker(tasks, engine_config=engine_config)
    results = []
    worker.file_finished.connect(results.append)
    loop = QEventLoop()
    worker.batch_finished.connect(lambda *a: loop.quit())
    worker.batch_cancelled.connect(lambda *a: loop.quit())
    worker.start()
    QTimer.singleShot(120000, loop.quit)
    loop.exec()
    while not worker.isFinished():
        app.processEvents()
    return results


def _isolation_check():
    ok = True
    d = tempfile.mkdtemp()
    png = Path(d) / "img.png"
    _make_png(png)
    entry = FileEntry(path=str(png), filename="img.png", extension=".png",
                      size=png.stat().st_size)
    cfg = _cfg(True, True, MockClient(exc=_mk_auth_error()))
    results = _run_worker([(0, entry)], cfg)
    ok &= _log("V6.6 worker 产出结果", len(results) == 1)
    r = results[0]
    ok &= _log("V6.6b AI 故障隔离: 文件仍 DONE", r.status == FileStatus.DONE,
               r.status.value)
    ok &= _log("V6.6c AI_PROVIDER_FAILURE warning 存在",
               any(w.code == "AI_PROVIDER_FAILURE" for w in r.warnings),
               [w.code for w in r.warnings])
    warn = next((w for w in r.warnings if w.code == "AI_PROVIDER_FAILURE"), None)
    ok &= _log("V6.6d warning 不含 key", warn is None or SECRET not in warn.message)

    report = ReportBuilder.build(r, entry)
    log = DiagnosticLogger(log_dir=d)
    log.record(report)
    line = Path(log.log_path).read_text(encoding="utf-8")
    ok &= _log("V6.6e 日志含 AI 故障且不含 key",
               "AI_PROVIDER_FAILURE" in line and SECRET not in line)
    return ok


# ---------------------------------------------------------------------------
# 7. Frozen-only checks
# ---------------------------------------------------------------------------
def _frozen_check():
    ok = True
    if getattr(sys, "frozen", False):
        for mod in ("markitdown", "markitdown_ocr", "openai", "src.ai_provider"):
            try:
                __import__(mod)
                ok &= _log(f"V6.7 冻结 EXE 内 {mod} 可导入", True)
            except Exception as exc:  # noqa: BLE001
                ok &= _log(f"V6.7 冻结 EXE 内 {mod} 可导入", False, repr(exc))
    else:
        _log("V6.7 源码模式跳过 frozen-only 检查", True)
    return ok


def main():
    app = QApplication.instance() or QApplication([])
    from src.main_window import MainWindow
    window = MainWindow()
    srv = _start_mock_server()
    try:
        results = [
            _no_ai_boot_check(window),
            _migration_check(),
            _provider_layer_check(),
            _connection_test_check(srv),
            _capability_toggle_check(),
            _isolation_check(),
            _frozen_check(),
        ]
    finally:
        srv.shutdown()
        srv.server_close()
    ok = all(results)
    print()
    print("ALL V0.6.0 RC CHECKS PASSED" if ok else "SOME V0.6.0 RC CHECKS FAILED")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
